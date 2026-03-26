"""
Resume parsing: PDF and DOCX text extraction with section-aware chunking.
"""

import re
import io
from typing import BinaryIO

import PyPDF2
import pdfplumber
from docx import Document


def parse_pdf(file_bytes: bytes) -> str:
    text = _parse_pdf_pypdf2(file_bytes)
    if not text or len(text.strip()) < 50:
        text = _parse_pdf_pdfplumber(file_bytes)
    return text.strip() if text else ""


def _parse_pdf_pypdf2(file_bytes: bytes) -> str:
    try:
        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        pages = []
        for page in reader.pages:
            t = page.extract_text()
            if t:
                pages.append(t)
        return "\n\n".join(pages)
    except Exception:
        return ""


def _parse_pdf_pdfplumber(file_bytes: bytes) -> str:
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            pages = []
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    pages.append(t)
            return "\n\n".join(pages)
    except Exception:
        return ""


def parse_docx(file_bytes: bytes) -> str:
    try:
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        return "\n".join(paragraphs)
    except Exception:
        return ""


def parse_resume(file_bytes: bytes, filename: str) -> str:
    ext = filename.lower().rsplit('.', 1)[-1] if '.' in filename else ''
    if ext == 'pdf':
        text = parse_pdf(file_bytes)
    elif ext in ('docx', 'doc'):
        text = parse_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file format: .{ext}. Only PDF and DOCX are supported.")

    if not text or len(text.strip()) < 20:
        raise ValueError(
            f"Could not extract meaningful text from '{filename}'. "
            "The file may be scanned/image-based or corrupted."
        )
    return text


SECTION_PATTERNS = [
    (r'(?i)^(work\s*experience|professional\s*experience|employment\s*history|experience)\s*$', 'experience'),
    (r'(?i)^(education|academic\s*background|qualifications|academic)\s*$', 'education'),
    (r'(?i)^(skills|technical\s*skills|core\s*competencies|technologies)\s*$', 'skills'),
    (r'(?i)^(projects|personal\s*projects|key\s*projects)\s*$', 'projects'),
    (r'(?i)^(certifications?|licenses?|credentials)\s*$', 'certifications'),
    (r'(?i)^(summary|objective|profile|about\s*me|professional\s*summary)\s*$', 'summary'),
    (r'(?i)^(awards?|achievements?|honors?)\s*$', 'awards'),
    (r'(?i)^(publications?|research|papers?)\s*$', 'publications'),
    (r'(?i)^(volunteer|community|extracurricular)\s*$', 'volunteer'),
    (r'(?i)^(competitive\s*programming\s*[&and]*\s*leadership|leadership\s*[&and]*\s*activities)\s*$', 'achievements'),
]

# Max words in a line for it to be considered a section header
_MAX_HEADER_WORDS = 8


def detect_section(text: str) -> str | None:
    """Detect if a line is a section header. Must be short and match a known pattern."""
    first_line = text.split('\n')[0].strip()
    # Section headers are short (≤ _MAX_HEADER_WORDS words)
    if len(first_line.split()) > _MAX_HEADER_WORDS:
        return None
    # Strip common formatting chars
    cleaned = re.sub(r'[:\-–—|•#*_=]', '', first_line).strip()
    if not cleaned:
        return None
    for pattern, section in SECTION_PATTERNS:
        if re.match(pattern, cleaned):
            return section
    return None


def _split_section_on_boundaries(section_text: str, chunk_size: int = 500) -> list[str]:
    """Split a large section on paragraph/bullet boundaries instead of raw word count."""
    # Split into paragraphs (double newline or bullet points)
    paragraphs = re.split(r'\n(?=\s*[•\-\*]|\n)', section_text)
    paragraphs = [p.strip() for p in paragraphs if p.strip()]

    chunks = []
    current_chunk = []
    current_words = 0

    for para in paragraphs:
        para_words = len(para.split())
        if current_words + para_words > chunk_size and current_chunk:
            chunks.append('\n'.join(current_chunk))
            current_chunk = [para]
            current_words = para_words
        else:
            current_chunk.append(para)
            current_words += para_words

    if current_chunk:
        chunks.append('\n'.join(current_chunk))

    return chunks


def chunk_text(text: str, chunk_size: int = 500, overlap: int = 100) -> list[dict]:
    lines = text.split('\n')
    raw_sections = []
    current_section = "general"
    current_lines = []

    for line in lines:
        detected = detect_section(line)
        if detected and current_lines:
            raw_sections.append((current_section, '\n'.join(current_lines)))
            current_section = detected
            current_lines = [line]
        else:
            if detected:
                current_section = detected
            current_lines.append(line)

    if current_lines:
        raw_sections.append((current_section, '\n'.join(current_lines)))

    # Merge orphan sections (header-only or very short) into the next section
    merged_sections = []
    i = 0
    while i < len(raw_sections):
        section_name, section_text = raw_sections[i]
        content = section_text.strip()
        # If section is just a header (< 50 chars of real content) and there's a next section
        content_without_header = '\n'.join(content.split('\n')[1:]).strip() if '\n' in content else ''
        if len(content_without_header) < 50 and i + 1 < len(raw_sections):
            # Merge: prepend this content to the next section, keep this section's name
            next_name, next_text = raw_sections[i + 1]
            merged_text = content + '\n' + next_text
            raw_sections[i + 1] = (section_name, merged_text)
        else:
            merged_sections.append((section_name, content))
        i += 1

    chunks = []
    chunk_index = 0

    for section_name, section_text in merged_sections:
        if not section_text.strip():
            continue

        # Prepend section label for embedding context
        labeled_text = f"[{section_name.title()}] {section_text.strip()}"

        words = labeled_text.split()
        if len(words) <= chunk_size:
            chunks.append({
                "chunk_index": chunk_index,
                "chunk_text": labeled_text,
                "section": section_name,
            })
            chunk_index += 1
        else:
            # Split on paragraph/bullet boundaries for semantic coherence
            sub_chunks = _split_section_on_boundaries(section_text, chunk_size)
            for sub in sub_chunks:
                labeled_sub = f"[{section_name.title()}] {sub.strip()}"
                if labeled_sub.strip():
                    chunks.append({
                        "chunk_index": chunk_index,
                        "chunk_text": labeled_sub,
                        "section": section_name,
                    })
                    chunk_index += 1

    if not chunks and text.strip():
        chunks.append({
            "chunk_index": 0,
            "chunk_text": text.strip()[:2000],
            "section": "general",
        })

    return chunks
